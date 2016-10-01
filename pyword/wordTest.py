# coding utf-8
# author xujipm


from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('文档标题11234223', 0)
p = document.add_paragraph('这是一个段落')
p.add_run('粗体').bold = True
p.add_run(' and some ')
p.add_run('斜体').italic = True
ru = p.add_run('ziticeshi','Emphasis')
ru.bold = True
ru.italic = True


document.add_heading('一111级标题', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')
pa = document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in ordered list', style='ListNumber')
pa.style = 'ListNumber'
#document.add_page_break()
document.add_picture('timg.jpg', width=Inches(2.25))
document.add_picture('timg.jpg', width=Inches(1.25))
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
row_cellls = table.add_row().cells
row_cellls[0].text = str(123)
row_cellls[1].text = str(23)
row_cellls[2].text = "item.Desc"
table.cell(1, 1).text = "new"
table.style = 'LightShading-Accent1'
document.add_paragraph("表格大小：" + str(len(table.rows)) +
                       "*" + str(len(table.columns)))
document.add_page_break()
document.save('范例.docx')
